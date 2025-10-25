# app.py
import os
import io
import time
import re
from datetime import datetime, date
import random
import pandas as pd
import streamlit as st

# =========================
# CONFIG
# =========================
# Personnel file path
PERSONNEL_PATH = os.path.join("input_csv", "personnel.xlsx")

# Project workbook path
PROJECT_WB_PATH = os.path.join("input_proj_excel", "employee.xlsx")
PROJECT_INFO_SHEET = "project_info"
EMPLOYEE_SHEET_NAME = "employee"  # where we'll write edited personnel

REQUIRED_COLS = ["Name", "Qualification", "Job Title", "From", "Years of Experience"]

SAVE_DIR = "temp_uploads"
OUTPUT_DOCX = os.path.abspath("Employees_CV.docx")
os.makedirs(SAVE_DIR, exist_ok=True)

st.set_page_config(page_title="Key Personnel • Editor & Bulk CVs", layout="wide")

# =========================
# STATE INIT
# =========================
def init_state():
    ss = st.session_state
    ss.setdefault("step", 1)                   # 1=Load & Review, 2=Roles, 3=Edit & Generate
    ss.setdefault("df_personnel", None)        # the working personnel dataframe
    ss.setdefault("df_project_info", None)     # loaded from PROJECT_WB_PATH / project_info
    ss.setdefault("roles", [])                 # list of dicts: {name, count, min_exp, degree_required}
    ss.setdefault("selection_mask", None)
    ss.setdefault("current_edit_path", None)   # temp_uploads file we keep overwriting after first save
    ss.setdefault("files_loaded", False)       # flag to track if files have been auto-loaded
    ss.setdefault("files_confirmed", False)    # flag to track if user confirmed the loaded files
    ss.setdefault("job_title_mode", None)      # "existing" or "assign_roles"
    ss.setdefault("defined_roles", [])         # roles defined for assignment in Step 3
    ss.setdefault("roles_defined_step3", False) # flag to track if roles are defined in step 3

init_state()

# =========================
# UTILS
# =========================
def ensure_required_cols(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    for c in REQUIRED_COLS:
        if c not in df.columns:
            df[c] = ""
    if "Assigned Role" not in df.columns:
        df["Assigned Role"] = ""
    # Ensure 'To' column exists with current date in MM-YYYY format
    if "To" not in df.columns:
        today = date.today()
        df["To"] = f"{today.month:02d}-{today.year}"
    return df

def save_temp_excel(df: pd.DataFrame, fixed_path: str | None = None) -> str:
    """
    If fixed_path is provided, overwrite it; else create a timestamped file in temp_uploads.
    Returns the path saved. Excludes 'Assigned Role' column from saved file.
    """
    # Create a copy to avoid modifying the original dataframe
    df_save = df.copy()
    
    # Remove 'Assigned Role' column from temp Excel files
    if "Assigned Role" in df_save.columns:
        df_save = df_save.drop(columns=["Assigned Role"])
    
    # Ensure date columns are stored as strings in MM-YYYY format
    if "From" in df_save.columns:
        df_save["From"] = df_save["From"].apply(lambda x: str(x) if pd.notna(x) and str(x).strip() != "" else "")
    if "To" in df_save.columns:
        # Store "To" as "Present" in Excel files
        df_save["To"] = "Present"
    
    # Ensure Years of Experience is integer
    if "Years of Experience" in df_save.columns:
        df_save["Years of Experience"] = df_save["Years of Experience"].apply(lambda x: int(float(x)) if pd.notna(x) else 0)
    
    if fixed_path:
        df_save.to_excel(fixed_path, index=False)
        return fixed_path
    ts = time.strftime("%Y%m%d_%H%M%S")
    p = os.path.join(SAVE_DIR, f"personnel_temp_{ts}.xlsx")
    df_save.to_excel(p, index=False)
    return p

def parse_from_to_date(val):
    """Parse date values from Excel, handling various formats and stripping time components"""
    if pd.isna(val) or val is None:
        return None

    # Handle datetime/Timestamp/date objects (strip time if present)
    if isinstance(val, (datetime, pd.Timestamp)):
        return date(val.year, val.month, 1)

    if isinstance(val, date):
        return date(val.year, val.month, 1)

    s = str(val).strip()
    if not s:
        return None

    # Handle MM-YYYY or MM/YYYY format
    mm_yyyy_match = re.match(r"^\s*(\d{1,2})[-/](\d{4})\s*$", s)
    if mm_yyyy_match:
        mm = max(1, min(12, int(mm_yyyy_match.group(1))))
        yy = int(mm_yyyy_match.group(2))
        return date(yy, mm, 1)

    # Handle year-only format (e.g., "2017" -> assumes 01-2017)
    if s.isdigit() and len(s) == 4:
        return date(int(s), 1, 1)
    
    # Handle DD-MM-YYYY or DD/MM/YYYY format (e.g., 01-01-2006, 09-12-2006)
    dd_mm_yyyy_match = re.match(r"^\s*(\d{1,2})[-/](\d{1,2})[-/](\d{4})\s*$", s)
    if dd_mm_yyyy_match:
        dd = int(dd_mm_yyyy_match.group(1))
        mm = max(1, min(12, int(dd_mm_yyyy_match.group(2))))
        yy = int(dd_mm_yyyy_match.group(3))
        return date(yy, mm, 1)

    # Try general date parsing (handles Excel dates with time)
    try:
        d = pd.to_datetime(s, errors="raise")
        return date(d.year, d.month, 1)
    except Exception:
        return None

def convert_to_mm_yyyy_format(val):
    """Convert various date formats to MM-YYYY string format"""
    if pd.isna(val) or val is None:
        return ""
    
    # Already a datetime/Timestamp/date object
    if isinstance(val, (datetime, pd.Timestamp, date)):
        return f"{val.month:02d}-{val.year}"
    
    s = str(val).strip()
    if not s:
        return ""
    
    # Already in MM-YYYY format
    if re.match(r"^\d{1,2}-\d{4}$", s):
        parts = s.split('-')
        mm = int(parts[0])
        yy = parts[1]
        return f"{mm:02d}-{yy}"
    
    # Year only (e.g., "2017") -> "01-2017"
    if s.isdigit() and len(s) == 4:
        return f"01-{s}"
    
    # DD-MM-YYYY format (e.g., "01-01-2006") -> "01-2006"
    dd_mm_yyyy_match = re.match(r"^(\d{1,2})[-/](\d{1,2})[-/](\d{4})$", s)
    if dd_mm_yyyy_match:
        mm = int(dd_mm_yyyy_match.group(2))
        yy = dd_mm_yyyy_match.group(3)
        return f"{mm:02d}-{yy}"
    
    # Try to parse as date
    try:
        d = pd.to_datetime(s, errors="raise")
        return f"{d.month:02d}-{d.year}"
    except:
        return ""  # Return empty instead of error

def years_since(d: date) -> int:
    """Calculate years of experience as integer (floor value, no decimals)"""
    if d is None:
        return 0
    today = date.today()
    end = date(today.year, today.month, 1)
    months = (end.year - d.year) * 12 + (end.month - d.month)
    # Return integer years (floor division)
    return int(months // 12)

def recalc_yoe_for_from_column(df: pd.DataFrame) -> pd.DataFrame:
    """Recalculate Years of Experience based on From date, returns integer years"""
    df = df.copy()
    y = []
    for _, r in df.iterrows():
        d = parse_from_to_date(r.get("From"))
        if d:
            y.append(years_since(d))
        else:
            # If no From date, try to use existing YOE as integer
            existing = r.get("Years of Experience", 0)
            try:
                y.append(int(float(existing)))
            except:
                y.append(0)
    df["Years of Experience"] = y
    return df

def ci_contains(text: str, needle: str) -> bool:
    t = str(text or "")
    n = str(needle or "")
    return n.lower() in t.lower() if n else False

def qualification_is_diploma(q: str) -> bool:
    return bool(re.search(r"\bdiploma\b", str(q or ""), flags=re.I))

def sync_job_title_with_assigned_role(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    if "Assigned Role" in df.columns:
        mask = df["Assigned Role"].astype(str).str.strip() != ""
        df.loc[mask, "Job Title"] = df.loc[mask, "Assigned Role"]
    return df

# ---- NEW: safe metric wrapper so React gets plain types ----
def safe_metric(label, value, delta=None):
    try:
        v = int(value)
    except Exception:
        try:
            v = int(float(value))
        except Exception:
            v = 0
    d = None if delta is None else str(delta)
    st.metric(str(label), v, delta=d)

# =========================
# AUTO-LOAD FILES
# =========================
def auto_load_files():
    """Auto-load personnel and project_info files on first run"""
    if st.session_state.files_loaded:
        return

    # Load Personnel file
    if os.path.exists(PERSONNEL_PATH):
        try:
            dfp = pd.read_excel(PERSONNEL_PATH)
            dfp = ensure_required_cols(dfp)
            dfp = recalc_yoe_for_from_column(dfp)
            # Ensure YOE is integer
            if "Years of Experience" in dfp.columns:
                dfp["Years of Experience"] = dfp["Years of Experience"].apply(lambda x: int(float(x)) if pd.notna(x) else 0)
            st.session_state.df_personnel = dfp
            st.session_state.personnel_load_status = f"✅ Loaded successfully: {len(dfp)} rows"
        except Exception as e:
            st.session_state.personnel_load_status = f"❌ Error loading file: {e}"
    else:
        st.session_state.personnel_load_status = f"❌ File not found: {PERSONNEL_PATH}"

    # Load Project Info file
    if os.path.exists(PROJECT_WB_PATH):
        try:
            xls = pd.ExcelFile(PROJECT_WB_PATH)
            if PROJECT_INFO_SHEET in xls.sheet_names:
                dfproj = pd.read_excel(xls, sheet_name=PROJECT_INFO_SHEET)
                st.session_state.df_project_info = dfproj
                st.session_state.project_load_status = f"✅ Loaded successfully: {len(dfproj)} rows"
            else:
                st.session_state.project_load_status = f"❌ Sheet '{PROJECT_INFO_SHEET}' not found"
        except Exception as e:
            st.session_state.project_load_status = f"❌ Error loading file: {e}"
    else:
        st.session_state.project_load_status = f"❌ File not found: {PROJECT_WB_PATH}"

    st.session_state.files_loaded = True

# =========================
# BULK GENERATOR (same structure you approved)
# =========================
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT_NAME = "Times New Roman"
FONT_SIZE = 8
LINE_SPACING = 1.15
INDENT_CM = 0.12

def to_dt_floor_month(val):
    """Convert any date value to a date object (first day of month), stripping time component"""
    if pd.isna(val) or val is None:
        return None

    # Handle date/datetime/Timestamp objects
    if isinstance(val, (datetime, pd.Timestamp)):
        return date(val.year, val.month, 1)

    if isinstance(val, date):
        return date(val.year, val.month, 1)

    s = str(val).strip()

    # Handle "Present" text
    if "present" in s.lower():
        today = date.today()
        return date(today.year, today.month, 1)

    # Handle year-only format
    if s.isdigit() and len(s) == 4:
        return date(int(s), 1, 1)

    # Handle MM-YYYY or MM/YYYY format
    mm_yyyy_match = re.match(r'^(\d{1,2})[-/](\d{4})$', s)
    if mm_yyyy_match:
        mm = int(mm_yyyy_match.group(1))
        yyyy = int(mm_yyyy_match.group(2))
        mm = max(1, min(12, mm))  # Ensure valid month
        return date(yyyy, mm, 1)

    # Try general date parsing
    try:
        d = pd.to_datetime(s, errors="raise")
        return date(d.year, d.month, 1)
    except Exception:
        return None

def format_mm_yyyy(val, allow_present=False):
    """Format any date value to MM-YYYY format for display in Word"""
    if pd.isna(val) or val is None:
        return ""

    # Handle 'Present' string
    s = str(val).strip()
    if allow_present and "present" in s.lower():
        return "Present"

    # Handle datetime/Timestamp objects (strip time component)
    if isinstance(val, (datetime, pd.Timestamp, date)):
        return val.strftime("%m-%Y")

    # Handle year-only format (e.g., "2020")
    if s.isdigit() and len(s) == 4:
        return f"01-{s}"

    # Handle MM-YYYY or similar formats already
    if re.match(r'^\d{1,2}[-/]\d{4}$', s):
        parts = re.split(r'[-/]', s)
        mm = parts[0].zfill(2)  # Ensure 2-digit month
        yyyy = parts[1]
        return f"{mm}-{yyyy}"

    # Try to parse as date and format
    try:
        d = pd.to_datetime(s, errors="raise")
        # Even if it has time component, only use date part
        return d.strftime("%m-%Y")
    except Exception:
        # If all else fails, return the original string
        return s

def bulletize(text):
    if not isinstance(text, str):
        return ""
    # Convert en-dash to regular dash
    text = text.replace("\u2013", "-").strip()
    # Remove leading/trailing double dashes specifically
    while text.startswith("--"):
        text = text[2:].strip()
    while text.endswith("--"):
        text = text[:-2].strip()
    # Split by double dash (--) to preserve single dashes in sentences
    raw = [t.strip() for t in text.split("--") if t.strip()]
    if not raw:
        return ""
    
    # Format each item: if it already starts with "-", keep it; otherwise add "- "
    formatted = []
    for item in raw:
        if item.startswith("-"):
            formatted.append(item)
        else:
            formatted.append("- " + item)
    
    return "\n".join(formatted)

def write_cell(cell, text, *, bold=False, font_name=FONT_NAME, font_size_pt=FONT_SIZE,
               indent_cm=INDENT_CM, line_spacing=LINE_SPACING):
    cell.text = ""
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.right_indent = Cm(indent_cm)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = line_spacing
    run = p.add_run(text if text is not None else "")
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bool(bold)

def add_row(table, text, bold=False):
    cell = table.add_row().cells[0]
    write_cell(cell, text, bold=bold)

def set_table_borders(table):
    """Set black borders for all cells in a table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Create table borders element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width (1/8 pt, so 4 = 0.5pt)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def run_bulk_generator(personnel_df: pd.DataFrame, project_info_df: pd.DataFrame | None, out_docx: str):
    import random as _random
    from docxcompose.composer import Composer
    df = personnel_df.copy()

    projects = pd.DataFrame()
    if project_info_df is not None and not project_info_df.empty:
        projects = project_info_df.rename(columns={
            "Start Date": "proj_start",
            "Work Completion date": "proj_end",
            "Company / Project / Position": "proj_cpp",
            "Relevant Technical & Managerial Experience": "proj_desc",
        }).copy()
        projects["proj_start_dt"] = projects["proj_start"].apply(to_dt_floor_month)
        projects["proj_end_dt"]   = projects["proj_end"].apply(to_dt_floor_month)

    # Track used projects to avoid duplicates
    used_project_indices = set()
    
    # Create list to store individual CV documents
    temp_docs = []

    for i, emp in df.iterrows():
        # Load a fresh copy of the template for each employee
        doc = Document('template/CV_template.docx')
        # Parse employee dates from personnel data
        emp_start_dt = to_dt_floor_month(emp.get("From"))
        emp_end_dt   = to_dt_floor_month(emp.get("To")) if "To" in df.columns else None
        
        # If To is not set or is "Present", use current date
        if emp_end_dt is None or str(emp.get("To", "")).lower() == "present":
            today = date.today()
            emp_end_dt = date(today.year, today.month, 1)

        # TABLE 1 – Personal Info
        table1 = doc.add_table(rows=0, cols=1)
        try:
            table1.style = "Table Grid"
        except KeyError:
            pass  # Style doesn't exist in template, use default
        set_table_borders(table1)

        add_row(table1, f"Position: {emp.get('Job Title','')}")
        add_row(table1, "Name of Bidder: Pioneer Foundation Engineers Private Limited")
        add_row(table1, f"Position: {emp.get('Job Title','')}")
        add_row(table1, "Personnel Information", bold=True)
        add_row(table1, f"Name: {emp.get('Name','')}")
        add_row(table1, f"Qualification / Certification / Licence / Training: {emp.get('Qualification','')}")
        add_row(table1, "Present Employment", bold=True)
        add_row(table1, "Name of Employer: Pioneer Foundation Engineers Private Limited")
        add_row(table1, "Address of Employer: Boomerang, B-2, 508/509, Off Chandivali Farm Rd, Chandivali, Powai, Mumbai, Maharashtra 400072")
        add_row(table1, "Telephone: 022 4801 1311")
        add_row(table1, "Contact (Manager / Personnel Officer): +91 99209 03578")
        add_row(table1, "Fax: –")
        add_row(table1, "E-mail: sales@pfepl.com")
        add_row(table1, f"Job Title: {emp.get('Job Title','')}")
        add_row(table1, f"Years with Present Employer: {emp.get('Years of Experience','')}")
        add_row(table1, "Mobile: +91 99209 03578")
        add_row(table1, "Professional Experience (Last 10 Years)", bold=True)

        gap_p = doc.add_paragraph("")
        gap_p.paragraph_format.space_before = Pt(0)
        gap_p.paragraph_format.space_after = Pt(2)

        # TABLE 2 – Experience Details (ONE project, random from eligible)
        table2 = doc.add_table(rows=1, cols=4)
        try:
            table2.style = "Table Grid"
        except KeyError:
            pass  # Style doesn't exist in template, use default
        set_table_borders(table2)

        hdr = table2.rows[0].cells
        write_cell(hdr[0], "From (MM-YYYY)", bold=True)
        write_cell(hdr[1], "To (MM-YYYY)", bold=True)
        write_cell(hdr[2], "Company / Project / Position", bold=True)
        write_cell(hdr[3], "Relevant Technical & Managerial Experience", bold=True)

        row = table2.add_row().cells
        chosen = None
        chosen_idx = None

        # Find eligible projects based on timeline overlap
        if not projects.empty:
            eligible = []
            for proj_idx, p in projects.iterrows():
                ps = p.get("proj_start_dt")  # Project start date
                pe = p.get("proj_end_dt")    # Project end date
                
                # Skip if both dates are missing
                if ps is None and pe is None:
                    continue
                
                # Handle missing project end date (ongoing project)
                if pe is None:
                    today = date.today()
                    pe = date(today.year, today.month, 1)
                
                # Handle missing project start date
                if ps is None:
                    ps = date(1900, 1, 1)
                
                # Check if project overlaps with employee tenure
                # Project is eligible if:
                # 1. Project ended AFTER employee started (pe >= emp_start_dt)
                # 2. Project started BEFORE employee ended (ps <= emp_end_dt)
                overlap = True
                if emp_start_dt and pe < emp_start_dt:
                    # Project completed before employee joined - NOT eligible
                    overlap = False
                if emp_end_dt and ps > emp_end_dt:
                    # Project started after employee left - NOT eligible
                    overlap = False
                
                if overlap:
                    eligible.append((proj_idx, p))
            
            # Try to pick an unused project first
            unused_eligible = [(idx, p) for idx, p in eligible if idx not in used_project_indices]
            
            if unused_eligible:
                # Pick random from unused projects
                chosen_idx, chosen = _random.choice(unused_eligible)
                used_project_indices.add(chosen_idx)
            elif eligible:
                # All projects used, pick random from any eligible
                chosen_idx, chosen = _random.choice(eligible)

        # Fill the experience row
        if chosen is not None:
            # Use EMPLOYEE's From and To dates (not project dates)
            from_disp = format_mm_yyyy(emp.get("From"))
            to_disp   = "Present"  # Always show Present for To
            
            # Use project's description
            cpp  = chosen.get("proj_cpp", "")
            desc = bulletize(chosen.get("proj_desc", ""))

            write_cell(row[0], from_disp)
            write_cell(row[1], to_disp)
            write_cell(row[2], str(cpp))
            write_cell(row[3], desc)
        else:
            # No eligible project found - use employee info only
            write_cell(row[0], format_mm_yyyy(emp.get("From")))
            write_cell(row[1], "Present")
            write_cell(row[2], f"Pioneer Foundation Engineers Pvt. Ltd. / {emp.get('Job Title','')}")
            write_cell(row[3], "")

        # Add page break after each CV (except the last one will be handled by merge)
        doc.add_page_break()

        # Save this individual CV temporarily
        temp_path = os.path.join(SAVE_DIR, f"temp_cv_{i}.docx")
        doc.save(temp_path)
        temp_docs.append(temp_path)

    # Merge all individual CVs into one final document
    if temp_docs:
        # Start with the first document
        master = Document(temp_docs[0])
        composer = Composer(master)
        
        # Append all other documents
        for temp_doc_path in temp_docs[1:]:
            doc_to_append = Document(temp_doc_path)
            composer.append(doc_to_append)
        
        # Save the final merged document
        composer.save(out_docx)
        
        # Clean up temporary files
        for temp_path in temp_docs:
            if os.path.exists(temp_path):
                os.remove(temp_path)

# =========================
# AUTO-LOAD FILES ON STARTUP
# =========================
auto_load_files()

# =========================
# LEFT NAV (vertical)
# =========================
st.sidebar.title("Navigation")
step = st.sidebar.radio(
    "Go to Step",
    options=[1, 2, 3],
    format_func=lambda x: {1: "Step 1 — Load & Review", 2: "Step 2 — Search (Optional)", 3: "Step 3 — Edit & Generate"}[x],
    index=st.session_state.step - 1,
    disabled=False
)

# Allow navigation to step 1 and 3 anytime, step 2 requires files confirmed
if step == 2 and not st.session_state.files_confirmed:
    st.warning("⚠️ Please confirm files in Step 1 before searching.")
    st.session_state.step = 1
elif step == 3 and not st.session_state.files_confirmed:
    st.warning("⚠️ Please confirm files in Step 1 before proceeding to Step 3.")
    st.session_state.step = 1
else:
    st.session_state.step = step

st.title("Key Personnel — Editor & Bulk CV Generator")

# =========================
# STEP 1 — REVIEW AUTO-LOADED FILES & UPLOAD
# =========================
if st.session_state.step == 1:
    st.header("Step 1 — Data Preparation & Validation")

    st.info("📂 Files are automatically loaded from the system. You can also upload your own files for processing.")

    # Upload Option Toggle
    use_upload = st.checkbox("📤 Upload Custom Files", value=False, help="Check this to upload your own personnel and project files")

    if not use_upload:
        # Use system files (auto-loaded)
        st.subheader("📋 System Files (Auto-Loaded)")
        
        # Display Personnel File Status and Content
        st.markdown("#### 1️⃣ Personnel File")
        col1, col2 = st.columns([3, 1])
        with col1:
            st.write(f"**File Path:** `{PERSONNEL_PATH}`")
            st.write(st.session_state.get("personnel_load_status", ""))
        with col2:
            if st.button("🔄 Reload Personnel", help="Reload the personnel file if you've made changes"):
                st.session_state.files_loaded = False
                st.session_state.files_confirmed = False
                st.session_state.df_personnel = None
                auto_load_files()
                st.rerun()

        if st.session_state.df_personnel is not None:
            # Validate and analyze the data
            df_check = st.session_state.df_personnel.copy()
            
            # Analysis
            with st.expander("📊 Data Analysis", expanded=False):
                issues = []
                warnings = []
                info_msgs = []
                
                # Check for missing required columns
                missing_cols = [col for col in REQUIRED_COLS if col not in df_check.columns]
                if missing_cols:
                    issues.append(f"❌ Missing required columns: {', '.join(missing_cols)}")
                
                # Check for missing values in key columns
                if "Name" in df_check.columns:
                    missing_names = df_check["Name"].isna().sum()
                    if missing_names > 0:
                        warnings.append(f"⚠️ {missing_names} rows have missing Names")
                
                if "From" in df_check.columns:
                    missing_from = df_check["From"].isna().sum()
                    if missing_from > 0:
                        warnings.append(f"⚠️ {missing_from} rows have missing 'From' dates - YOE will be 0")
                    
                    # Check for various date formats
                    year_only = 0
                    dd_mm_yyyy = 0
                    other_formats = 0
                    for val in df_check["From"].dropna():
                        if pd.notna(val):
                            s = str(val).strip()
                            if s.isdigit() and len(s) == 4:
                                year_only += 1
                            elif re.match(r"^\d{1,2}[-/]\d{1,2}[-/]\d{4}$", s):
                                dd_mm_yyyy += 1
                            elif not re.match(r"^\d{1,2}[-/]\d{4}$", s) and not isinstance(val, (datetime, pd.Timestamp, date)):
                                other_formats += 1
                    
                    if year_only > 0 or dd_mm_yyyy > 0 or other_formats > 0:
                        info_msgs.append(f"ℹ️ Date formats detected: {year_only} year-only (2017), {dd_mm_yyyy} DD-MM-YYYY (01-01-2006), {other_formats} other - will auto-convert to MM-YYYY in Step 3")
                
                if "To" in df_check.columns:
                    missing_to = df_check["To"].isna().sum()
                    if missing_to > 0:
                        st.info(f"ℹ️ {missing_to} rows have missing 'To' dates - will be set to 'Present' in Step 3")
                
                if "Years of Experience" in df_check.columns:
                    missing_yoe = df_check["Years of Experience"].isna().sum()
                    if missing_yoe > 0:
                        st.info(f"ℹ️ {missing_yoe} rows have missing 'Years of Experience' - will be auto-calculated in Step 3")
                
                if issues:
                    for issue in issues:
                        st.error(issue)
                if warnings:
                    for warning in warnings:
                        st.warning(warning)
                if info_msgs:
                    for info in info_msgs:
                        st.info(info)
                
                if not issues and not warnings:
                    st.success("✅ All data looks good!")
                
                st.caption(f"📈 Total rows: {len(df_check)} | Columns: {len(df_check.columns)}")
            
            st.dataframe(st.session_state.df_personnel, use_container_width=True, height=300)
            st.caption(f"Total rows: {len(st.session_state.df_personnel)}")
        else:
            st.warning("Personnel file could not be loaded. Please check the file path and try reloading.")

        st.divider()

        # Display Project Info File Status and Content
        st.markdown("#### 2️⃣ Project Info File")
        st.write(f"**File Path:** `{PROJECT_WB_PATH}`")
        st.write(f"**Sheet Name:** `{PROJECT_INFO_SHEET}`")
        st.write(st.session_state.get("project_load_status", ""))

        if st.session_state.df_project_info is not None:
            st.dataframe(st.session_state.df_project_info, use_container_width=True, height=300)
            st.caption(f"Total rows: {len(st.session_state.df_project_info)}")
        else:
            st.warning("Project info file could not be loaded. Please check the file path.")

    else:
        # Upload custom files
        st.subheader("📤 Upload Custom Personnel File")
        st.info("ℹ️ **Note:** Project info will always be loaded from the system file. Only personnel data can be uploaded.")
        
        # Personnel File Upload
        st.markdown("#### 📤 Upload Personnel File")
        up_personnel_file = st.file_uploader("Upload Personnel Excel File (.xlsx)", type=["xlsx"], accept_multiple_files=False, key="custom_personnel_uploader")
        
        if up_personnel_file is not None:
            try:
                # Read all sheets
                xls_personnel = pd.ExcelFile(up_personnel_file)
                sheet_names = xls_personnel.sheet_names
                
                st.success(f"✅ File uploaded: {up_personnel_file.name}")
                st.info(f"📑 Available sheets: {', '.join(sheet_names)}")
                
                # Sheet selection
                selected_personnel_sheet = st.selectbox(
                    "Select Personnel Sheet",
                    options=sheet_names,
                    key="personnel_sheet_select",
                    help="Choose the sheet containing personnel data"
                )
                
                if selected_personnel_sheet:
                    dfp = pd.read_excel(xls_personnel, sheet_name=selected_personnel_sheet)
                    
                    # Validation and Analysis
                    st.markdown("##### 📊 Data Validation & Analysis")
                    
                    issues = []
                    warnings = []
                    info_msgs = []
                    
                    # Check for required columns
                    missing_cols = [col for col in REQUIRED_COLS if col not in dfp.columns]
                    if missing_cols:
                        issues.append(f"❌ Missing required columns: {', '.join(missing_cols)}")
                    
                    # Check data quality
                    if "Name" in dfp.columns:
                        missing_names = dfp["Name"].isna().sum()
                        if missing_names > 0:
                            warnings.append(f"⚠️ {missing_names} rows have missing Names")
                    
                    if "From" in dfp.columns:
                        missing_from = dfp["From"].isna().sum()
                        if missing_from > 0:
                            info_msgs.append(f"ℹ️ {missing_from} rows have missing 'From' dates - will default to 0 YOE")
                        
                        # Check for various date formats
                        year_only = 0
                        dd_mm_yyyy = 0
                        other_formats = 0
                        for val in dfp["From"].dropna():
                            if pd.notna(val):
                                s = str(val).strip()
                                if s.isdigit() and len(s) == 4:
                                    year_only += 1
                                elif re.match(r"^\d{1,2}[-/]\d{1,2}[-/]\d{4}$", s):
                                    dd_mm_yyyy += 1
                                elif not re.match(r"^\d{1,2}[-/]\d{4}$", s) and not isinstance(val, (datetime, pd.Timestamp, date)):
                                    other_formats += 1
                        
                        if year_only > 0 or dd_mm_yyyy > 0 or other_formats > 0:
                            info_msgs.append(f"ℹ️ Date formats detected: {year_only} year-only (2017), {dd_mm_yyyy} DD-MM-YYYY (01-01-2006), {other_formats} other - will auto-convert to MM-YYYY in Step 3")
                    
                    if "To" not in dfp.columns or dfp["To"].isna().all():
                        info_msgs.append("ℹ️ 'To' column missing or empty - will be set to 'Present' in Step 3")
                    
                    if "Years of Experience" not in dfp.columns or dfp["Years of Experience"].isna().all():
                        info_msgs.append("ℹ️ 'Years of Experience' missing - will be auto-calculated from 'From' date in Step 3")
                    
                    # Display validation results
                    for issue in issues:
                        st.error(issue)
                    for warning in warnings:
                        st.warning(warning)
                    for info in info_msgs:
                        st.info(info)
                    
                    if not issues:
                        st.success(f"✅ Data validation passed! Ready to process {len(dfp)} rows")
                        
                        # Process and store the data
                        dfp = ensure_required_cols(dfp)
                        dfp = recalc_yoe_for_from_column(dfp)
                        # Ensure YOE is integer
                        if "Years of Experience" in dfp.columns:
                            dfp["Years of Experience"] = dfp["Years of Experience"].apply(lambda x: int(float(x)) if pd.notna(x) else 0)
                        
                        st.session_state.df_personnel = dfp
                        st.session_state.current_edit_path = None
                        
                        # Show preview
                        st.dataframe(dfp, use_container_width=True, height=300)
                        st.caption(f"Total rows: {len(dfp)} | Columns: {len(dfp.columns)}")
                    else:
                        st.error("❌ Cannot proceed with this file. Please fix the issues above.")
                        
            except Exception as e:
                st.error(f"❌ Failed to read uploaded file: {e}")
                import traceback
                with st.expander("Error Details"):
                    st.code(traceback.format_exc())
        
        st.divider()
        
        # Show system Project Info (read-only)
        st.markdown("#### 📋 Project Info (System File)")
        st.write(f"**File Path:** `{PROJECT_WB_PATH}`")
        st.write(f"**Sheet Name:** `{PROJECT_INFO_SHEET}`")
        st.write(st.session_state.get("project_load_status", ""))
        
        if st.session_state.df_project_info is not None:
            st.dataframe(st.session_state.df_project_info, use_container_width=True, height=300)
            st.caption(f"Total projects: {len(st.session_state.df_project_info)}")
        else:
            st.warning("Project info file could not be loaded from system.")

    st.divider()

    # Confirm Button
    st.subheader("✅ Confirm and Proceed")

    can_proceed = (st.session_state.df_personnel is not None and
                   st.session_state.df_project_info is not None)

    if can_proceed:
        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if st.button("✔️ Confirm & Search Personnel", type="primary", use_container_width=True):
                st.session_state.files_confirmed = True
                st.session_state.step = 2
                st.success("Files confirmed! Moving to Step 2...")
                st.rerun()
        with col2:
            if st.button("⏩ Skip Search, Go to Edit", use_container_width=True):
                st.session_state.files_confirmed = True
                st.session_state.step = 3
                st.success("Files confirmed! Skipping search, moving to Step 3...")
                st.rerun()
        with col3:
            st.caption("💡 You can search for roles in Step 2, or skip directly to editing in Step 3")
    else:
        st.error("⚠️ Cannot proceed. Please ensure both Personnel and Project Info files are loaded successfully.")

# =========================
# STEP 2 — ENHANCED ROLE DEFINITION & SEARCH
# =========================
if st.session_state.step == 2 and st.session_state.df_personnel is not None:
    st.header("Step 2 — Define Roles & Search Personnel")

    with st.expander("➕ Add / Update Role", expanded=True):
        st.markdown("### Role Configuration")

        # Row 1: Role Name and Required Count
        col1, col2 = st.columns([3, 1])
        with col1:
            r_name = st.text_input(
                "Role Name *",
                placeholder="e.g., Civil Engineer, Project Manager",
                help="Enter the role/position name. This will be matched against Job Title."
            )
        with col2:
            r_count = st.number_input(
                "Required Count *",
                min_value=0,
                step=1,
                value=1,
                help="How many people you need for this role"
            )

        # Row 2: Keywords for Qualification Search
        st.markdown("---")
        st.markdown("#### Qualification Filtering (Optional)")

        r_keywords = st.text_input(
            "Qualification Keywords",
            placeholder="e.g., civil, mechanical, electrical (separate multiple keywords with commas)",
            help="Enter keywords to search in Qualifications. Multiple keywords work with OR logic (matches ANY keyword, not all required). Example: 'civil, mechanical' will match people with EITHER civil OR mechanical qualifications. Leave empty to skip qualification filtering."
        )

        col3, col4 = st.columns(2)
        with col3:
            search_mode = st.radio(
                "Search Mode",
                options=["Contains (Anywhere)", "Exact Word Match"],
                index=0,
                help="Contains: Finds keywords anywhere in qualifications (e.g., 'civil' matches 'B.E. Civil', 'Civil Engineering')\n\nExact Word: Matches complete word only (e.g., 'civil' matches 'B.E. Civil' but not 'civilization')"
            )
        with col4:
            include_diploma = st.checkbox(
                "Include Diploma Certified",
                value=False,
                help="If CHECKED: Include people with 'Diploma' in qualifications\nIf UNCHECKED: Only include degree holders (exclude Diploma)"
            )

        # Row 3: Experience Filter
        st.markdown("---")
        st.markdown("#### Experience Filtering (Optional)")

        col5, col6 = st.columns([1, 3])
        with col5:
            r_min_exp = st.number_input(
                "Minimum Years of Experience",
                min_value=0.0,
                step=0.5,
                value=0.0,
                format="%.1f",
                help="Minimum years of experience required. Set to 0 to skip experience filtering."
            )
        with col6:
            st.info("💡 **Tips:**\n- Leave keywords empty to search all qualifications\n- Multiple keywords use OR logic: 'civil, mechanical' matches EITHER civil OR mechanical\n- Keywords are case-insensitive\n- Example: 'civil, mechanical' will match 'B.E. Civil' (has civil) AND 'B.Tech. Mechanical' (has mechanical)")

        # Save and Clear buttons
        st.markdown("---")
        colb1, colb2, colb3 = st.columns([2, 2, 4])
        with colb1:
            if st.button("💾 Save Role", type="primary", use_container_width=True):
                try:
                    name = (r_name or "").strip()
                    if not name:
                        st.error("❌ Role Name is required!")
                    elif r_count <= 0:
                        st.error("❌ Required Count must be greater than 0!")
                    else:
                        # Process keywords
                        keywords_list = []
                        if r_keywords and r_keywords.strip():
                            keywords_list = [k.strip().lower() for k in r_keywords.split(',') if k.strip()]

                        role_data = {
                            "name": name,
                            "count": int(r_count),
                            "min_exp": float(r_min_exp),
                            "keywords": keywords_list,
                            "search_mode": "exact" if "Exact" in search_mode else "contains",
                            "include_diploma": bool(include_diploma)
                        }

                        # Upsert role
                        found = False
                        for i, rr in enumerate(st.session_state.roles):
                            if rr["name"].lower() == name.lower():
                                st.session_state.roles[i] = role_data
                                found = True
                                break
                        if not found:
                            st.session_state.roles.append(role_data)

                        st.success(f"✅ Role saved: {name}")
                        st.rerun()
                except Exception as e:
                    st.error(f"❌ Failed to save role: Wrong data or invalid input.")
                    st.error(f"Error details: {str(e)}")

        with colb2:
            if st.button("🗑️ Clear All Roles", use_container_width=True):
                st.session_state.roles = []
                st.success("All roles cleared.")
                st.rerun()

    # Display defined roles
    if st.session_state.roles:
        st.markdown("### 📋 Defined Roles")
        roles_display = []
        for r in st.session_state.roles:
            roles_display.append({
                "Role": r["name"],
                "Count": r["count"],
                "Min Experience": r["min_exp"],
                "Keywords": ", ".join(r.get("keywords", [])) if r.get("keywords") else "None",
                "Search Mode": r.get("search_mode", "loose").title(),
                "Include Diploma": "Yes" if r.get("include_diploma", False) else "No"
            })
        st.dataframe(pd.DataFrame(roles_display), use_container_width=True)
    else:
        st.info("ℹ️ No roles defined yet. Add at least one role to enable search.")

    # Search functionality
    st.divider()
    if st.session_state.roles:
        if st.button("🔎 Start Search", type="primary", use_container_width=False):
            try:
                # Verify data is available
                if st.session_state.df_personnel is None or st.session_state.df_personnel.empty:
                    st.error("❌ No personnel data found. Please load data in Step 1 first.")
                    st.stop()
                
                df = st.session_state.df_personnel.copy()
                
                # Verify required columns exist
                required_search_cols = ["Name", "Qualification", "Job Title", "Years of Experience"]
                missing_cols = [col for col in required_search_cols if col not in df.columns]
                if missing_cols:
                    st.error(f"❌ Required columns missing: {', '.join(missing_cols)}. Please check your data.")
                    st.stop()
                
                df["__yoe__"] = pd.to_numeric(df["Years of Experience"], errors="coerce").fillna(0.0)

                summary_rows = []

                for role in st.session_state.roles:
                    role_name = role["name"]
                    min_exp = float(role.get("min_exp", 0.0))
                    required = int(role["count"])
                    keywords = role.get("keywords", [])
                    search_mode = role.get("search_mode", "contains")
                    include_diploma = role.get("include_diploma", False)

                    # Title matching
                    title_hit = df["Job Title"].apply(lambda x: ci_contains(x, role_name))

                    # Experience filter
                    if min_exp > 0:
                        exp_ok = df["__yoe__"] >= min_exp
                    else:
                        exp_ok = pd.Series([True] * len(df), index=df.index)

                    # Qualification filter
                    if keywords:
                        if search_mode == "exact":
                            # Exact word match (case-insensitive)
                            # Matches "civil" in "B.E. Civil", "Civil Engineering", etc.
                            def exact_match(qual_text):
                                if pd.isna(qual_text):
                                    return False
                                # Split by common delimiters and clean up
                                qual_lower = str(qual_text).lower()
                                # Split by spaces, dots, commas, slashes, parentheses
                                words = re.split(r'[\s\.\,\/\(\)\-]+', qual_lower)
                                words = [w.strip() for w in words if w.strip()]
                                # Check if any keyword matches any word exactly
                                return any(kw in words for kw in keywords)

                            qual_ok = df["Qualification"].apply(exact_match)
                        else:
                            # Contains: keyword found anywhere (substring match)
                            def contains_match(qual_text):
                                if pd.isna(qual_text):
                                    return False
                                qual_lower = str(qual_text).lower()
                                return any(kw in qual_lower for kw in keywords)

                            qual_ok = df["Qualification"].apply(contains_match)
                    else:
                        qual_ok = pd.Series([True] * len(df), index=df.index)

                    # Diploma filter
                    if include_diploma:
                        # Include everyone (diploma or not)
                        diploma_ok = pd.Series([True] * len(df), index=df.index)
                    else:
                        # Exclude diploma holders
                        diploma_ok = ~df["Qualification"].apply(qualification_is_diploma)

                    # Combined filters - Enhanced categorization
                    # Adjust categorization based on which filters are actually active
                    has_qual_filter = bool(keywords)
                    has_exp_filter = min_exp > 0
                    
                    if has_qual_filter and has_exp_filter:
                        # All filters active - full categorization
                        # 1. Fully matched: Title + Experience + Qualification all match
                        fully_mask = title_hit & exp_ok & qual_ok & diploma_ok
                        
                        # 2. Title + Qualification match, but Experience insufficient
                        title_qual_no_exp_mask = title_hit & qual_ok & diploma_ok & (~exp_ok)
                        
                        # 3. Qualification + Experience match, but Title mismatch
                        qual_exp_no_title_mask = (~title_hit) & exp_ok & qual_ok & diploma_ok
                        
                        # 4. Qualification only (no title, no experience)
                        qual_only_mask = qual_ok & diploma_ok & (~exp_ok) & (~title_hit)
                        
                        # 5. No match at all
                        no_match_mask = ~(fully_mask | title_qual_no_exp_mask | qual_exp_no_title_mask | qual_only_mask)
                    
                    elif has_qual_filter:
                        # Only qualification filter active
                        # 1. Fully matched: Title + Qualification match
                        fully_mask = title_hit & qual_ok & diploma_ok
                        
                        # 2. Qualification only (no title match)
                        qual_only_mask = (~title_hit) & qual_ok & diploma_ok
                        
                        # 3. No qualification match
                        no_match_mask = ~qual_ok | ~diploma_ok
                        
                        # Not applicable
                        title_qual_no_exp_mask = pd.Series([False] * len(df), index=df.index)
                        qual_exp_no_title_mask = pd.Series([False] * len(df), index=df.index)
                    
                    elif has_exp_filter:
                        # Only experience filter active
                        # 1. Fully matched: Title + Experience match
                        fully_mask = title_hit & exp_ok
                        
                        # 2. Experience but wrong title
                        qual_exp_no_title_mask = (~title_hit) & exp_ok
                        
                        # 3. No experience match
                        no_match_mask = ~exp_ok
                        
                        # Not applicable
                        title_qual_no_exp_mask = pd.Series([False] * len(df), index=df.index)
                        qual_only_mask = pd.Series([False] * len(df), index=df.index)
                    
                    else:
                        # No filters - only title matching
                        # 1. Fully matched: Title matches
                        fully_mask = title_hit
                        
                        # 2. No title match
                        no_match_mask = ~title_hit
                        
                        # Not applicable
                        title_qual_no_exp_mask = pd.Series([False] * len(df), index=df.index)
                        qual_exp_no_title_mask = pd.Series([False] * len(df), index=df.index)
                        qual_only_mask = pd.Series([False] * len(df), index=df.index)

                    # Extract results - just get the data, no formatting
                    def safe_extract(mask, df):
                        """Safely extract dataframe slice"""
                        if mask.sum() == 0:
                            return pd.DataFrame(columns=["Name", "Qualification", "Job Title", "From", "Years of Experience"])
                        result = df[mask][["Name", "Qualification", "Job Title", "From", "Years of Experience"]].copy()
                        result = result.reset_index(drop=True)
                        return result

                    fully = safe_extract(fully_mask, df)
                    title_qual_no_exp = safe_extract(title_qual_no_exp_mask, df)
                    qual_exp_no_title = safe_extract(qual_exp_no_title_mask, df)
                    qual_only = safe_extract(qual_only_mask, df)
                    no_match = safe_extract(no_match_mask, df)

                    missing = max(required - len(fully), 0)
                    
                    # Build summary row with only relevant columns based on active filters
                    summary_row = {
                        "Role": role_name,
                        "Required": required,
                        "Fully Matched": len(fully),
                    }
                    
                    # Add columns based on which filters are active
                    if has_qual_filter and has_exp_filter:
                        summary_row["Title+Qual (No Exp)"] = len(title_qual_no_exp)
                        summary_row["Qual+Exp (No Title)"] = len(qual_exp_no_title)
                        summary_row["Qualification Only"] = len(qual_only)
                    elif has_qual_filter:
                        summary_row["Qualification Only"] = len(qual_only)
                    elif has_exp_filter:
                        summary_row["Exp (No Title)"] = len(qual_exp_no_title)
                    
                    summary_row["No Match"] = len(no_match)
                    summary_row["Missing"] = missing
                    
                    summary_rows.append(summary_row)

                    with st.expander(f"📌 {role_name} — Search Results", expanded=True):
                        # Display search criteria used
                        criteria_parts = []
                        criteria_parts.append(f"Keywords: {', '.join(keywords) if keywords else 'None'}")
                        if keywords:
                            criteria_parts.append(f"Mode: {search_mode.title()}")
                        if min_exp > 0:
                            criteria_parts.append(f"Min Exp: {min_exp} yrs")
                        criteria_parts.append(f"Include Diploma: {'Yes' if include_diploma else 'No'}")
                        st.caption(f"🔍 Search Criteria: {' | '.join(criteria_parts)}")

                        # Show metrics based on active filters
                        if has_qual_filter and has_exp_filter:
                            c1, c2, c3 = st.columns(3)
                            with c1:
                                safe_metric("✅ Fully Matched",
                                            len(fully),
                                            f"-{missing} shortage" if missing > 0 else ("Target met!" if required > 0 else None))
                            with c2:
                                safe_metric("⚠️ Title+Qual (Low Exp)", len(title_qual_no_exp))
                            with c3:
                                safe_metric("🔍 Qual + Exp (Job Title Mismatch)", len(qual_exp_no_title))

                            c4, c5 = st.columns(2)
                            with c4:
                                safe_metric("📋 Qualification Only", len(qual_only))
                            with c5:
                                safe_metric("❌ No Match", len(no_match))
                        elif has_qual_filter:
                            c1, c2, c3 = st.columns(3)
                            with c1:
                                safe_metric("✅ Fully Matched",
                                            len(fully),
                                            f"-{missing} shortage" if missing > 0 else ("Target met!" if required > 0 else None))
                            with c2:
                                safe_metric("📋 Qualification Only", len(qual_only))
                            with c3:
                                safe_metric("❌ No Match", len(no_match))
                        elif has_exp_filter:
                            c1, c2, c3 = st.columns(3)
                            with c1:
                                safe_metric("✅ Fully Matched",
                                            len(fully),
                                            f"-{missing} shortage" if missing > 0 else ("Target met!" if required > 0 else None))
                            with c2:
                                safe_metric("👤 Exp (No Title)", len(qual_exp_no_title))
                            with c3:
                                safe_metric("❌ No Match", len(no_match))
                        else:
                            c1, c2 = st.columns(2)
                            with c1:
                                safe_metric("✅ Fully Matched",
                                            len(fully),
                                            f"-{missing} shortage" if missing > 0 else ("Target met!" if required > 0 else None))
                            with c2:
                                safe_metric("❌ No Match", len(no_match))

                        st.markdown("---")

                        # Always show fully matched
                        st.markdown("**✅ Fully Matched**")
                        if len(fully) > 0:
                            try:
                                fully_display = fully.copy()
                                for col in fully_display.columns:
                                    fully_display[col] = fully_display[col].apply(lambda x: str(x) if pd.notna(x) else "")
                                st.dataframe(fully_display, use_container_width=True)
                            except Exception:
                                st.warning("⚠️ Table display issue. Showing simplified view:")
                                for idx, row in fully.iterrows():
                                    st.text(f"• {row['Name']} - {row['Qualification']} - {row['Job Title']} - {row.get('From', '')} - {row['Years of Experience']} yrs")
                        else:
                            st.info("ℹ️ No fully matched candidates found.")

                        # Show Title+Qual (Low Exp) only if both qual and exp filters are set
                        if has_qual_filter and has_exp_filter and len(title_qual_no_exp) > 0:
                            st.markdown("**⚠️ Title + Qualification Match (Experience Insufficient)**")
                            try:
                                title_qual_display = title_qual_no_exp.copy()
                                for col in title_qual_display.columns:
                                    title_qual_display[col] = title_qual_display[col].apply(lambda x: str(x) if pd.notna(x) else "")
                                st.dataframe(title_qual_display, use_container_width=True)
                                st.caption("💡 These candidates have the right position and qualifications but need more experience.")
                            except Exception:
                                st.warning("⚠️ Table display issue. Showing simplified view:")
                                for idx, row in title_qual_no_exp.iterrows():
                                    st.text(f"• {row['Name']} - {row['Qualification']} - {row['Job Title']} - {row.get('From', '')} - {row['Years of Experience']} yrs")

                        # Show Qual+Exp (Wrong Title) only if both qual and exp filters are set
                        if has_qual_filter and has_exp_filter and len(qual_exp_no_title) > 0:
                            st.markdown("**🔍 Qualification + Experience Match (Title Mismatch)**")
                            try:
                                qual_exp_display = qual_exp_no_title.copy()
                                for col in qual_exp_display.columns:
                                    qual_exp_display[col] = qual_exp_display[col].apply(lambda x: str(x) if pd.notna(x) else "")
                                st.dataframe(qual_exp_display, use_container_width=True)
                                st.caption("💡 These candidates have the right qualifications and experience but different job title.")
                            except Exception:
                                st.warning("⚠️ Table display issue. Showing simplified view:")
                                for idx, row in qual_exp_no_title.iterrows():
                                    st.text(f"• {row['Name']} - {row['Qualification']} - {row['Job Title']} - {row.get('From', '')} - {row['Years of Experience']} yrs")

                        # Show Qualification Only if qual filter is set
                        if has_qual_filter and len(qual_only) > 0:
                            st.markdown("**📋 Qualification Only**")
                            try:
                                qual_only_display = qual_only.copy()
                                for col in qual_only_display.columns:
                                    qual_only_display[col] = qual_only_display[col].apply(lambda x: str(x) if pd.notna(x) else "")
                                st.dataframe(qual_only_display, use_container_width=True)
                                if has_exp_filter:
                                    st.caption("💡 These candidates have the right qualifications but title and/or experience don't match.")
                                else:
                                    st.caption("💡 These candidates have the right qualifications but title doesn't match.")
                            except Exception:
                                st.warning("⚠️ Table display issue. Showing simplified view:")
                                for idx, row in qual_only.iterrows():
                                    st.text(f"• {row['Name']} - {row['Qualification']} - {row['Job Title']} - {row.get('From', '')} - {row['Years of Experience']} yrs")
                        
                        # Show Experience (No Title) if only exp filter is set
                        if has_exp_filter and not has_qual_filter and len(qual_exp_no_title) > 0:
                            st.markdown("**👤 Has Experience (Wrong Title)**")
                            try:
                                exp_display = qual_exp_no_title.copy()
                                for col in exp_display.columns:
                                    exp_display[col] = exp_display[col].apply(lambda x: str(x) if pd.notna(x) else "")
                                st.dataframe(exp_display, use_container_width=True)
                                st.caption("💡 These candidates have the required experience but different job title.")
                            except Exception:
                                st.warning("⚠️ Table display issue. Showing simplified view:")
                                for idx, row in qual_exp_no_title.iterrows():
                                    st.text(f"• {row['Name']} - {row['Qualification']} - {row['Job Title']} - {row.get('From', '')} - {row['Years of Experience']} yrs")

                st.markdown("### 📊 Summary by Role")
                try:
                    summary_df = pd.DataFrame(summary_rows)
                    
                    # Fill NaN values with "N/A" and format numbers as integers
                    for col in summary_df.columns:
                        if col != "Role":  # Don't modify the Role column
                            # Convert to numeric, fill NaN with "N/A", then format as integer where applicable
                            summary_df[col] = summary_df[col].apply(lambda x: "N/A" if pd.isna(x) else (int(x) if isinstance(x, (int, float)) else x))
                    
                    st.dataframe(summary_df, use_container_width=True)
                except Exception:
                    st.warning("⚠️ Could not display summary table. Showing as text:")
                    for row in summary_rows:
                        st.text(f"{row['Role']}: {row['Fully Matched']}/{row['Required']} matched ({row['Missing']} missing)")
            
            except Exception as e:
                st.error("❌ Search failed: Wrong data format or no valid data found.")
                st.error(f"Error details: {str(e)}")
                with st.expander("🔍 Technical Details"):
                    import traceback
                    st.code(traceback.format_exc())

    st.divider()
    st.info("✅ Review the search results above, then proceed to **Step 3** to edit and generate CVs.")

# =========================
# STEP 3 — INLINE EDIT + PERSIST + GENERATE
# =========================
if st.session_state.step == 3 and st.session_state.df_personnel is not None:
    st.header("Step 3 — Edit Personnel & Generate CVs")

    # =========================
    # MODE SELECTION - Use Existing Job Titles or Assign New Roles
    # =========================
    if st.session_state.job_title_mode is None:
        st.subheader("🔧 Job Title Configuration")
        st.info("Choose how you want to handle Job Titles for personnel")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📋 Use Existing Job Titles", type="primary", use_container_width=True):
                st.session_state.job_title_mode = "existing"
                st.rerun()
            st.caption("Keep current job titles from the loaded Excel file. Edit them as needed.")
        
        with col2:
            if st.button("🎯 Assign New Roles", type="primary", use_container_width=True):
                st.session_state.job_title_mode = "assign_roles"
                # Clear job titles when switching to assign roles mode
                df_work = st.session_state.df_personnel.copy()
                df_work["Job Title"] = ""
                st.session_state.df_personnel = df_work
                st.rerun()
            st.caption("Clear all job titles and assign them from a list of defined roles.")
        
        st.stop()  # Stop here until mode is selected
    
    # Show current mode with option to reset
    col_mode, col_reset = st.columns([3, 1])
    with col_mode:
        mode_display = "📋 Using Existing Job Titles" if st.session_state.job_title_mode == "existing" else "🎯 Assigning New Roles"
        st.info(f"**Current Mode:** {mode_display}")
    with col_reset:
        if st.button("🔄 Change Mode", use_container_width=True):
            st.session_state.job_title_mode = None
            st.session_state.defined_roles = []
            st.session_state.roles_defined_step3 = False
            st.rerun()
    
    st.divider()
    
    # =========================
    # ROLE DEFINITION (if Assign New Roles mode)
    # =========================
    if st.session_state.job_title_mode == "assign_roles":
        if not st.session_state.roles_defined_step3:
            st.subheader("🎯 Define Job Roles")
            st.warning("⚠️ Please define at least one role before proceeding to edit personnel.")
            
            with st.expander("➕ Add Role", expanded=True):
                st.markdown("### Role Definition")
                
                role_name_input = st.text_input(
                    "Role Name *",
                    placeholder="e.g., Civil Engineer, Project Manager, Site Supervisor",
                    help="Enter the role/position name that will be assigned to personnel",
                    key="step3_role_name"
                )
                
                col_btn1, col_btn2, col_btn3 = st.columns([2, 2, 4])
                with col_btn1:
                    if st.button("💾 Add Role", type="primary", use_container_width=True):
                        if not role_name_input.strip():
                            st.error("❌ Role name cannot be empty!")
                        elif role_name_input.strip() in st.session_state.defined_roles:
                            st.warning("⚠️ This role already exists!")
                        else:
                            st.session_state.defined_roles.append(role_name_input.strip())
                            st.success(f"✅ Added role: {role_name_input.strip()}")
                            st.rerun()
                
                with col_btn2:
                    if st.button("🗑️ Clear All Roles", use_container_width=True):
                        st.session_state.defined_roles = []
                        st.success("All roles cleared.")
                        st.rerun()
            
            # Display defined roles
            if st.session_state.defined_roles:
                st.markdown("### 📋 Defined Roles")
                roles_display_df = pd.DataFrame({
                    "Role Name": st.session_state.defined_roles,
                    "Actions": ["✅ Active"] * len(st.session_state.defined_roles)
                })
                st.dataframe(roles_display_df, use_container_width=True, hide_index=True)
                
                # Remove individual role
                st.markdown("##### Remove Role")
                col_select, col_remove = st.columns([3, 1])
                with col_select:
                    role_to_remove = st.selectbox(
                        "Select role to remove",
                        options=st.session_state.defined_roles,
                        key="role_remove_select"
                    )
                with col_remove:
                    st.write("")
                    st.write("")
                    if st.button("🗑️ Remove", use_container_width=True):
                        st.session_state.defined_roles.remove(role_to_remove)
                        st.success(f"Removed role: {role_to_remove}")
                        st.rerun()
                
                st.divider()
                
                # Confirm and proceed
                if st.button("✅ Confirm Roles & Proceed to Edit", type="primary", use_container_width=False):
                    st.session_state.roles_defined_step3 = True
                    st.success(f"✅ {len(st.session_state.defined_roles)} role(s) confirmed! You can now edit personnel.")
                    st.rerun()
            else:
                st.info("ℹ️ No roles defined yet. Add at least one role to continue.")
            
            st.stop()  # Stop here until roles are defined
    
    # Quick guide
    with st.expander("📖 Quick Guide", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            **🔧 Editing Options:**
            - **Editable Fields**: Name, Qualification, Job Title, From (MM-YYYY)
            - **Auto-Calculated**: To (always "Present"), Years of Experience
            - **Job Title Mode**: Choose between existing titles or assign new roles
            - **Add Row**: Click "+" button or use form below (all fields mandatory)
            - **From Date Format**: MM-YYYY (e.g., 01-2020, 11-2022)
            - **Delete Row**: Select row(s) with checkbox and use delete button
            - **Empty Rows**: Automatically removed on save
            """)
        with col2:
            st.markdown("""
            **⚡ Quick Actions:**
            - **Bulk Job Title**: Select rows → Choose role/enter title → Apply
            - **Bulk From Date**: Select rows → Enter date (MM-YYYY) → Apply
            - **Add Single User**: All fields mandatory, YOE auto-calculated
            - **Auto-Save**: Changes save automatically to temp Excel
            - **YOE**: Calculated in whole years (integer, no decimals)
            - **Warnings**: Format issues shown as warnings, not errors
            """)
    
    if st.session_state.job_title_mode == "assign_roles":
        st.info("💡 **Assign Roles Mode**: Job titles selected from defined roles. Use dropdown in table or bulk tool. 'Custom' option available for non-standard titles.")
    else:
        st.info("💡 **Features**: Edit Name, Qualification, Job Title & From date inline. Add rows with +. Bulk assign job titles/dates. Auto-save on changes. Empty rows auto-removed. Format warnings won't block saving.")

    # PRE-PROCESSING: Analyze and convert date formats
    df_work = st.session_state.df_personnel.copy()
    today = date.today()
    default_to = f"{today.month:02d}-{today.year}"
    
    # Analyze "From" column for various formats
    format_issues = []
    converted_dates = []
    
    if "From" in df_work.columns:
        for idx, val in enumerate(df_work["From"]):
            if pd.notna(val) and str(val).strip() != "":
                original = str(val).strip()
                converted = convert_to_mm_yyyy_format(val)
                
                if converted == "":
                    # Could not convert
                    format_issues.append(f"Row {idx + 1}: '{original}' - unsupported format")
                elif converted != original:
                    # Was converted from different format
                    converted_dates.append(f"Row {idx + 1}: '{original}' → '{converted}'")
                    df_work.at[idx, "From"] = converted
    
    # Show format analysis if any conversions or issues found
    if format_issues or converted_dates:
        with st.expander("📋 Date Format Analysis & Conversion", expanded=True):
            if converted_dates:
                st.success(f"✅ Auto-converted {len(converted_dates)} date(s) to MM-YYYY format:")
                st.caption("**Conversions Applied:**")
                for msg in converted_dates[:10]:  # Show first 10
                    st.write(f"• {msg}")
                if len(converted_dates) > 10:
                    st.caption(f"... and {len(converted_dates) - 10} more")
                st.info("💡 **Supported formats**: Year only (2017 → 01-2017), DD-MM-YYYY (01-01-2006 → 01-2006), MM-YYYY (06-2022 → 06-2022)")
            
            if format_issues:
                st.warning(f"⚠️ {len(format_issues)} date(s) have unsupported formats and were set to empty:")
                for issue in format_issues[:10]:
                    st.write(f"• {issue}")
                if len(format_issues) > 10:
                    st.caption(f"... and {len(format_issues) - 10} more")
                st.error("❌ **Action Required**: Please check and correct the formats. Use MM-YYYY format (e.g., 01-2020, 06-2022)")
    
    # Ensure "To" column exists and is formatted
    if "To" not in df_work.columns:
        df_work["To"] = default_to
    else:
        df_work["To"] = df_work["To"].apply(lambda x: convert_to_mm_yyyy_format(x) if pd.notna(x) and str(x).strip() != "" else default_to)
    
    # Apply conversions to "From" column
    if "From" in df_work.columns:
        df_work["From"] = df_work["From"].apply(lambda x: convert_to_mm_yyyy_format(x) if pd.notna(x) else "")

    st.session_state.df_personnel = df_work

    # --- COLLAPSIBLE FORM FOR ADDING SINGLE USER ---
    with st.expander("➕ Add Single User", expanded=False):
        st.markdown("### Add New Personnel")
        st.info("ℹ️ **Note**: All fields are mandatory. 'From' date must be in MM-YYYY format (e.g., 01-2020). Years of Experience and To date will be calculated automatically.")
        
        col1, col2 = st.columns(2)
        with col1:
            new_name = st.text_input("Name *", key="new_name", placeholder="Enter full name")
            new_qual = st.text_input("Qualification *", key="new_qual", placeholder="e.g., B.E. Civil")
            
            # Job Title - dropdown if in assign_roles mode, text input otherwise
            if st.session_state.job_title_mode == "assign_roles" and st.session_state.defined_roles:
                job_title_options = st.session_state.defined_roles + ["Custom"]
                selected_job_title = st.selectbox(
                    "Job Title *",
                    options=job_title_options,
                    help="Select from defined roles or choose Custom",
                    key="new_job_title_select"
                )
                if selected_job_title == "Custom":
                    new_job_title = st.text_input(
                        "Enter Custom Job Title *",
                        key="new_job_title_custom",
                        placeholder="e.g., Site Engineer"
                    )
                else:
                    new_job_title = selected_job_title
            else:
                new_job_title = st.text_input("Job Title *", key="new_job_title", placeholder="e.g., Site Engineer")
        with col2:
            new_from = st.text_input("From (MM-YYYY) *", key="new_from", placeholder="e.g., 01-2020 or 11-2022")
            st.caption("📅 To: **Present** (auto-set)")
            st.caption("📊 Years of Experience: **Auto-calculated** based on From date")

        if st.button("✅ Add User", type="primary", use_container_width=True):
            # Validation
            errors = []
            if not new_name.strip():
                errors.append("❌ Name is required")
            if not new_qual.strip():
                errors.append("❌ Qualification is required")
            if not new_job_title.strip():
                errors.append("❌ Job Title is required")
            if not new_from.strip():
                errors.append("❌ From date is required")
            elif not re.match(r'^\d{1,2}-\d{4}$', new_from.strip()):
                errors.append("❌ From date must be in MM-YYYY format (e.g., 01-2020 or 11-2022)")
            
            if errors:
                for error in errors:
                    st.error(error)
            else:
                # Parse and validate the From date
                try:
                    parts = new_from.strip().split('-')
                    month = int(parts[0])
                    year = int(parts[1])
                    if month < 1 or month > 12:
                        st.error("❌ Invalid month. Must be between 01 and 12")
                    elif year < 1900 or year > 2100:
                        st.error("❌ Invalid year")
                    else:
                        # Format to ensure 2-digit month
                        formatted_from = f"{month:02d}-{year}"
                        
                        # Calculate YOE automatically
                        from_date = parse_from_to_date(formatted_from)
                        calculated_yoe = years_since(from_date) if from_date else 0
                        
                        new_row = {
                            "Name": new_name.strip(),
                            "Qualification": new_qual.strip(),
                            "Job Title": new_job_title.strip(),
                            "From": formatted_from,
                            "To": "Present",
                            "Years of Experience": calculated_yoe
                        }
                        df_work = st.session_state.df_personnel.copy()
                        df_work = pd.concat([df_work, pd.DataFrame([new_row])], ignore_index=True)
                        st.session_state.df_personnel = df_work
                        
                        # Auto-save
                        if st.session_state.current_edit_path is None:
                            st.session_state.current_edit_path = save_temp_excel(df_work)
                        else:
                            save_temp_excel(df_work, fixed_path=st.session_state.current_edit_path)
                        
                        st.success(f"✅ Added user: {new_name.strip()} | YOE: {calculated_yoe} years")
                        st.rerun()
                except Exception as e:
                    st.error(f"❌ Invalid date format: {e}")

    st.divider()

    # Show current personnel count and job title assignment status
    col_info1, col_info2, col_info3, col_info4 = st.columns([1, 1, 1, 2])
    with col_info1:
        st.metric("Total Personnel", len(st.session_state.df_personnel))
    with col_info2:
        selected_count = 0
        if st.session_state.selection_mask is not None and st.session_state.selection_mask.any():
            selected_count = int(st.session_state.selection_mask.sum())
        st.metric("Selected", selected_count)
    with col_info3:
        # Show job title assignment status
        df_check = st.session_state.df_personnel.copy()
        if "Job Title" in df_check.columns:
            assigned = (~df_check["Job Title"].isna()) & (df_check["Job Title"].astype(str).str.strip() != "")
            assigned_count = assigned.sum()
            unassigned_count = len(df_check) - assigned_count
            
            if unassigned_count > 0:
                st.metric("Job Titles Assigned", f"{assigned_count}/{len(df_check)}", delta=f"-{unassigned_count} missing", delta_color="inverse")
            else:
                st.metric("Job Titles Assigned", f"{assigned_count}/{len(df_check)}", delta="All assigned")
    with col_info4:
        st.caption("✏️ **Editable Table**: Edit inline. + to add rows. Empty rows auto-deleted. Select checkboxes for bulk operations.")

    # Show quick alert if there are any missing critical fields
    df_alert_check = st.session_state.df_personnel.copy()
    missing_fields = []
    
    # Check Name
    if "Name" in df_alert_check.columns:
        missing_names = (df_alert_check["Name"].isna() | (df_alert_check["Name"].astype(str).str.strip() == "")).sum()
        if missing_names > 0:
            missing_fields.append(f"{missing_names} Name(s)")
    
    # Check Job Title
    if "Job Title" in df_alert_check.columns:
        missing_jt = (df_alert_check["Job Title"].isna() | (df_alert_check["Job Title"].astype(str).str.strip() == "")).sum()
        if missing_jt > 0:
            missing_fields.append(f"{missing_jt} Job Title(s)")
    
    # Check Qualification
    if "Qualification" in df_alert_check.columns:
        missing_qual = (df_alert_check["Qualification"].isna() | (df_alert_check["Qualification"].astype(str).str.strip() == "")).sum()
        if missing_qual > 0:
            missing_fields.append(f"{missing_qual} Qualification(s)")
    
    # Check From date
    if "From" in df_alert_check.columns:
        missing_from = (df_alert_check["From"].isna() | (df_alert_check["From"].astype(str).str.strip() == "")).sum()
        if missing_from > 0:
            missing_fields.append(f"{missing_from} From Date(s)")
    
    if missing_fields:
        st.warning(f"⚠️ **Missing Required Data**: {', '.join(missing_fields)} - Please complete all required fields before downloading or generating CVs.")

    # --- Editable grid with auto-save ---
    df_edit = st.session_state.df_personnel.copy()
    
    # Remove 'Assigned Role' from display (not shown in UI or temp Excel)
    if "Assigned Role" in df_edit.columns:
        df_edit = df_edit.drop(columns=["Assigned Role"])
    
    # Convert date columns to string and ensure proper format
    if "From" in df_edit.columns:
        df_edit["From"] = df_edit["From"].apply(lambda x: str(x) if pd.notna(x) and str(x).strip() != "" else "")
    
    # Display "Present" for To column
    if "To" in df_edit.columns:
        df_edit["To"] = "Present"
    
    # Ensure YOE is integer
    if "Years of Experience" in df_edit.columns:
        df_edit["Years of Experience"] = df_edit["Years of Experience"].apply(lambda x: int(float(x)) if pd.notna(x) else 0)
    
    if "Select" not in df_edit.columns:
        df_edit.insert(0, "Select", False)

    # Column configuration - Name, Qualification, Job Title, From are editable
    # Job Title uses selectbox if in assign_roles mode
    if st.session_state.job_title_mode == "assign_roles" and st.session_state.defined_roles:
        column_config = {
            "Select": st.column_config.CheckboxColumn("Select", default=False, width="small"),
            "Name": st.column_config.TextColumn("Name", width="medium", required=True),
            "Qualification": st.column_config.TextColumn("Qualification", width="medium", required=True),
            "Job Title": st.column_config.SelectboxColumn(
                "Job Title",
                width="medium",
                options=st.session_state.defined_roles,
                required=True,
                help="Select from defined roles"
            ),
            "From": st.column_config.TextColumn("From (MM-YYYY)", width="small", required=True, help="Format: MM-YYYY (e.g., 01-2020)"),
            "To": st.column_config.TextColumn("To", width="small"),
            "Years of Experience": st.column_config.NumberColumn("YOE (Years)", width="small", format="%d"),
        }
    else:
        column_config = {
            "Select": st.column_config.CheckboxColumn("Select", default=False, width="small"),
            "Name": st.column_config.TextColumn("Name", width="medium", required=True),
            "Qualification": st.column_config.TextColumn("Qualification", width="medium", required=True),
            "Job Title": st.column_config.TextColumn("Job Title", width="medium", required=True),
            "From": st.column_config.TextColumn("From (MM-YYYY)", width="small", required=True, help="Format: MM-YYYY (e.g., 01-2020)"),
            "To": st.column_config.TextColumn("To", width="small"),
            "Years of Experience": st.column_config.NumberColumn("YOE (Years)", width="small", format="%d"),
        }
    
    # Disable columns that shouldn't be editable
    disabled_columns = ["To", "Years of Experience"]

    if st.session_state.job_title_mode == "assign_roles" and st.session_state.defined_roles:
        st.caption("⚠️ **Editable**: Name, Qualification, Job Title (dropdown: defined roles), From (MM-YYYY format) | **Auto-calculated**: To (Present), Years of Experience")
    else:
        st.caption("⚠️ **Editable**: Name, Qualification, Job Title, From (MM-YYYY format) | **Auto-calculated**: To (Present), Years of Experience")
    
    edited = st.data_editor(
        df_edit,
        use_container_width=True,
        num_rows="dynamic",
        key="data_editor_step3",
        column_config=column_config,
        disabled=disabled_columns,
        hide_index=True,
    )

    # Commit edits with auto-save and validation
    if isinstance(edited, pd.DataFrame):
        if "Select" in edited.columns:
            st.session_state.selection_mask = edited["Select"].fillna(False).astype(bool).values
            edited = edited.drop(columns=["Select"])
        else:
            st.session_state.selection_mask = None

        # Filter out completely empty rows before validation
        def is_row_empty(row):
            """Check if a row is completely empty (all values are NaN or empty strings)"""
            for val in row:
                if pd.notna(val) and str(val).strip() != "":
                    return False
            return True
        
        # Remove completely empty rows
        empty_mask = edited.apply(is_row_empty, axis=1)
        if empty_mask.any():
            num_empty = empty_mask.sum()
            edited = edited[~empty_mask].reset_index(drop=True)
            st.info(f"ℹ️ Removed {num_empty} empty row(s) automatically")

        # Validate From column format (MM-YYYY) - only for non-empty values
        validation_warnings = []
        if "From" in edited.columns and len(edited) > 0:
            for idx, row in edited.iterrows():
                val = row.get("From")
                # Only validate if From has a value
                if pd.notna(val) and str(val).strip() != "":
                    val_str = str(val).strip()
                    # Check if it's already in datetime format (from Excel)
                    if not re.match(r'^\d{1,2}[-/]\d{4}$', val_str):
                        # Try to parse as datetime
                        try:
                            pd.to_datetime(val)
                            # It's a valid date, we'll convert it in recalc_yoe_for_from_column
                        except:
                            validation_warnings.append(f"Row {idx + 1}: 'From' date format may need correction (use MM-YYYY like 01-2020)")
                    else:
                        # Validate month range for MM-YYYY format
                        try:
                            parts = val_str.split('-') if '-' in val_str else val_str.split('/')
                            month = int(parts[0])
                            if month < 1 or month > 12:
                                validation_warnings.append(f"Row {idx + 1}: Invalid month '{month}'. Must be between 01 and 12")
                        except:
                            validation_warnings.append(f"Row {idx + 1}: Could not parse 'From' date")
        
        # Show validation warnings but don't block saving
        if validation_warnings:
            with st.expander("⚠️ Data Format Warnings", expanded=True):
                st.warning("Some rows have format issues. Please review and correct:")
                for warning in validation_warnings:
                    st.write(f"• {warning}")
                st.info("💡 **Tip**: Use MM-YYYY format (e.g., 01-2020, 11-2022). Empty values are okay and will be handled automatically.")
        
        # Always save, even with warnings
        # Recalculate YOE based on From dates
        edited = recalc_yoe_for_from_column(edited)
        
        # Add back 'Assigned Role' column if it doesn't exist (for internal use)
        if "Assigned Role" not in edited.columns:
            edited["Assigned Role"] = ""
        
        # Ensure To column is "Present"
        edited["To"] = "Present"
        
        # Ensure required columns exist
        edited = ensure_required_cols(edited)
        
        st.session_state.df_personnel = edited

        # Auto-save to temp file (first time create, then overwrite)
        if st.session_state.current_edit_path is None:
            st.session_state.current_edit_path = save_temp_excel(edited)  # timestamped
        else:
            save_temp_excel(edited, fixed_path=st.session_state.current_edit_path)

    st.divider()
    st.subheader("🛠️ Bulk Assignment")
    st.caption("Select rows in the table above, choose a field to update, then apply to all selected rows")

    # Unified bulk assignment interface
    col_select, col_value, col_button = st.columns([1, 2, 1])
    
    with col_select:
        bulk_column = st.selectbox(
            "Select Field to Update",
            options=["Job Title", "From", "Qualification"],
            key="bulk_column_select"
        )
    
    with col_value:
        # Dynamic help text and input based on selected column and mode
        if bulk_column == "From":
            placeholder_text = "e.g., 01-2020, 11-2022"
            help_text = "📅 Format: MM-YYYY (e.g., 01-2020). Years of Experience will auto-calculate."
            bulk_value = st.text_input(
                "Value to Assign",
                placeholder=placeholder_text,
                help=help_text,
                key="bulk_value_input"
            )
            st.caption(help_text)
        elif bulk_column == "Qualification":
            placeholder_text = "e.g., B.E. Civil, M.Tech. Mechanical, Diploma Civil"
            help_text = "🎓 Enter degree in short form like: B.E. Civil, M.Tech. Civil, Diploma Civil"
            bulk_value = st.text_input(
                "Value to Assign",
                placeholder=placeholder_text,
                help=help_text,
                key="bulk_value_input"
            )
            st.caption(help_text)
        else:  # Job Title
            # If in assign_roles mode, show dropdown with defined roles + Custom option
            if st.session_state.job_title_mode == "assign_roles" and st.session_state.defined_roles:
                help_text = "💼 Select from defined roles or choose 'Custom' to enter your own"
                role_options = st.session_state.defined_roles + ["Custom"]
                selected_role = st.selectbox(
                    "Select Role to Assign",
                    options=role_options,
                    help=help_text,
                    key="bulk_role_select"
                )
                st.caption(help_text)
                
                # If Custom is selected, show text input
                if selected_role == "Custom":
                    bulk_value = st.text_input(
                        "Enter Custom Job Title",
                        placeholder="e.g., Site Engineer, Project Manager",
                        help="💼 Enter custom job title",
                        key="bulk_custom_value"
                    )
                else:
                    bulk_value = selected_role
            else:
                # Use text input for existing mode
                placeholder_text = "e.g., Site Engineer, Project Manager"
                help_text = "💼 Enter full job title as required. Cannot be empty."
                bulk_value = st.text_input(
                    "Value to Assign",
                    placeholder=placeholder_text,
                    help=help_text,
                    key="bulk_value_input"
                )
                st.caption(help_text)
    
    with col_button:
        st.write("")  # Spacing
        st.write("")  # Spacing
        if st.button("✅ Apply to Selected", type="primary", use_container_width=True):
            mask = st.session_state.selection_mask
            if mask is None or not mask.any():
                st.warning("⚠️ No rows selected. Please select rows using checkboxes.")
            elif not bulk_value.strip():
                st.error(f"❌ {bulk_column} cannot be empty!")
            else:
                # Validation based on column type
                if bulk_column == "From":
                    if not re.match(r'^\d{1,2}-\d{4}$', bulk_value.strip()):
                        st.error("❌ Invalid format. Please use MM-YYYY (e.g., 01-2020)")
                    else:
                        # Validate month range
                        try:
                            parts = bulk_value.strip().split('-')
                            month = int(parts[0])
                            if month < 1 or month > 12:
                                st.error(f"❌ Invalid month '{month}'. Must be between 01 and 12")
                            else:
                                dfx = st.session_state.df_personnel.copy()
                                dfx.loc[mask, "From"] = bulk_value.strip()
                                dfx = recalc_yoe_for_from_column(dfx)
                                st.session_state.df_personnel = dfx
                                if st.session_state.current_edit_path is None:
                                    st.session_state.current_edit_path = save_temp_excel(dfx)
                                else:
                                    save_temp_excel(dfx, fixed_path=st.session_state.current_edit_path)
                                st.success(f"✅ Assigned '{bulk_value}' to {int(mask.sum())} person(s). YOE recalculated.")
                                st.rerun()
                        except:
                            st.error("❌ Invalid date format")
                else:
                    # Job Title or Qualification
                    dfx = st.session_state.df_personnel.copy()
                    dfx.loc[mask, bulk_column] = bulk_value.strip()
                    if bulk_column == "Job Title" and "Assigned Role" in dfx.columns:
                        dfx.loc[mask, "Assigned Role"] = bulk_value.strip()
                    st.session_state.df_personnel = dfx
                    if st.session_state.current_edit_path is None:
                        st.session_state.current_edit_path = save_temp_excel(dfx)
                    else:
                        save_temp_excel(dfx, fixed_path=st.session_state.current_edit_path)
                    st.success(f"✅ Assigned '{bulk_value}' to {bulk_column} for {int(mask.sum())} person(s)")
                    st.rerun()

    st.divider()
    
    # Additional operations
    col_delete, col_selected, col_info = st.columns([1, 1, 2])
    
    with col_delete:
        if st.button("🗑️ Delete Selected", key="delete_btn"):
            mask = st.session_state.selection_mask
            if mask is None or not mask.any():
                st.warning("⚠️ No rows selected. Use checkboxes to select rows to delete.")
            else:
                num_to_delete = int(mask.sum())
                kept = st.session_state.df_personnel.loc[~mask].reset_index(drop=True)
                st.session_state.df_personnel = kept
                if st.session_state.current_edit_path is None:
                    st.session_state.current_edit_path = save_temp_excel(kept)
                else:
                    save_temp_excel(kept, fixed_path=st.session_state.current_edit_path)
                st.session_state.selection_mask = None
                st.success(f"✅ Deleted {num_to_delete} row(s)")
                st.rerun()
    
    with col_selected:
        mask = st.session_state.selection_mask
        selected_count = int(mask.sum()) if mask is not None and mask.any() else 0
        st.info(f"📊 **Selected**: {selected_count} row(s)")
    
    with col_info:
        st.caption("💡 **Tip**: Select rows with checkboxes, then delete or use bulk tools above")

    # Show auto-save status
    if st.session_state.current_edit_path:
        with st.expander("💾 Auto-Save Status", expanded=False):
            st.success(f"✅ Changes are being auto-saved to: `{st.session_state.current_edit_path}`")
            st.caption("Every edit you make is automatically saved. The file is updated whenever you:")
            st.caption("• Edit cells in the table above")
            st.caption("• Add a new user via the form")
            st.caption("• Use bulk assignment tools")
            st.caption("• Delete rows")

    st.divider()
    st.subheader("📥 Download & Generate")
    
    # =========================
    # VALIDATION BEFORE DOWNLOAD/GENERATE
    # =========================
    validation_errors = []
    validation_warnings = []
    df_validation = st.session_state.df_personnel.copy()
    
    # Helper function to get name or row number
    def get_row_identifier(idx):
        if "Name" in df_validation.columns and pd.notna(df_validation.loc[idx, "Name"]) and str(df_validation.loc[idx, "Name"]).strip() != "":
            return f"Row {idx+1}: {df_validation.loc[idx, 'Name']}"
        return f"Row {idx+1}"
    
    # Check for empty Names - CRITICAL
    if "Name" in df_validation.columns:
        empty_names = df_validation["Name"].isna() | (df_validation["Name"].astype(str).str.strip() == "")
        if empty_names.any():
            empty_count = empty_names.sum()
            empty_indices = df_validation[empty_names].index.tolist()
            validation_errors.append({
                "title": "❌ Missing Names",
                "count": empty_count,
                "details": [f"Row {idx+1}" for idx in empty_indices[:10]],
                "message": f"{empty_count} personnel have no name assigned"
            })
    
    # Check for empty Job Titles - CRITICAL
    if "Job Title" in df_validation.columns:
        empty_job_titles = df_validation["Job Title"].isna() | (df_validation["Job Title"].astype(str).str.strip() == "")
        if empty_job_titles.any():
            empty_indices = df_validation[empty_job_titles].index.tolist()
            empty_names = []
            for idx in empty_indices:
                empty_names.append(get_row_identifier(idx))
            validation_errors.append({
                "title": "❌ Missing Job Titles",
                "count": len(empty_indices),
                "details": empty_names[:10],
                "message": f"{len(empty_indices)} personnel have no job title assigned"
            })
    
    # Check for empty Qualifications - CRITICAL
    if "Qualification" in df_validation.columns:
        empty_qual = df_validation["Qualification"].isna() | (df_validation["Qualification"].astype(str).str.strip() == "")
        if empty_qual.any():
            empty_count = empty_qual.sum()
            empty_indices = df_validation[empty_qual].index.tolist()
            qual_names = []
            for idx in empty_indices:
                qual_names.append(get_row_identifier(idx))
            validation_errors.append({
                "title": "❌ Missing Qualifications",
                "count": empty_count,
                "details": qual_names[:10],
                "message": f"{empty_count} personnel have no qualification assigned"
            })
    
    # Check for empty From dates - CRITICAL
    if "From" in df_validation.columns:
        empty_from = df_validation["From"].isna() | (df_validation["From"].astype(str).str.strip() == "")
        if empty_from.any():
            empty_count = empty_from.sum()
            empty_indices = df_validation[empty_from].index.tolist()
            from_names = []
            for idx in empty_indices:
                from_names.append(get_row_identifier(idx))
            validation_errors.append({
                "title": "❌ Missing From Dates",
                "count": empty_count,
                "details": from_names[:10],
                "message": f"{empty_count} personnel have no 'From' date (Years of Experience will be 0)"
            })
    
    # Check for invalid From date format - WARNING
    if "From" in df_validation.columns:
        invalid_format = []
        for idx, row in df_validation.iterrows():
            val = row.get("From")
            if pd.notna(val) and str(val).strip() != "":
                val_str = str(val).strip()
                # Check if format is MM-YYYY or valid date
                if not re.match(r'^\d{1,2}[-/]\d{4}$', val_str):
                    try:
                        pd.to_datetime(val)
                    except:
                        invalid_format.append((idx, val_str))
        
        if invalid_format:
            details = []
            for idx, val in invalid_format[:10]:
                details.append(f"{get_row_identifier(idx)} - '{val}'")
            validation_warnings.append({
                "title": "⚠️ Invalid From Date Format",
                "count": len(invalid_format),
                "details": details,
                "message": f"{len(invalid_format)} personnel have 'From' dates in unsupported format (use MM-YYYY like 01-2020)"
            })
    
    # Check for Zero Years of Experience - WARNING
    if "Years of Experience" in df_validation.columns:
        zero_yoe = df_validation["Years of Experience"].fillna(0) == 0
        if zero_yoe.any():
            zero_count = zero_yoe.sum()
            zero_indices = df_validation[zero_yoe].index.tolist()
            yoe_names = []
            for idx in zero_indices[:10]:
                yoe_names.append(get_row_identifier(idx))
            validation_warnings.append({
                "title": "⚠️ Zero Years of Experience",
                "count": zero_count,
                "details": yoe_names,
                "message": f"{zero_count} personnel have 0 years of experience (may need 'From' date correction)"
            })
    
    # Display validation results in expander
    if validation_errors or validation_warnings:
        with st.expander("🔍 Data Validation Results - Click to view details", expanded=True):
            # Summary at the top
            col_summary1, col_summary2 = st.columns(2)
            with col_summary1:
                if validation_errors:
                    st.error(f"**❌ {len(validation_errors)} Critical Issue(s)**")
                else:
                    st.success("**✅ No Critical Issues**")
            with col_summary2:
                if validation_warnings:
                    st.warning(f"**⚠️ {len(validation_warnings)} Warning(s)**")
                else:
                    st.success("**✅ No Warnings**")
            
            st.markdown("---")
            
            # Show critical errors first
            if validation_errors:
                st.error("**❌ CRITICAL ISSUES - Must be fixed before proceeding:**")
                st.caption("These issues will prevent you from downloading Excel or generating CVs")
                st.markdown("")
                
                for i, error in enumerate(validation_errors, 1):
                    st.markdown(f"**{i}. {error['title']}**")
                    st.error(error['message'])
                    
                    # Show affected rows
                    with st.container():
                        st.markdown(f"**📋 Affected rows ({error['count']} total):**")
                        for detail in error['details']:
                            st.text(f"  • {detail}")
                        if error['count'] > 10:
                            st.caption(f"  ... and {error['count'] - 10} more rows")
                    st.markdown("")
                
                st.markdown("---")
            
            # Show warnings
            if validation_warnings:
                st.warning("**⚠️ WARNINGS - Recommended to fix:**")
                st.caption("These won't block download/generation but should be addressed for data quality")
                st.markdown("")
                
                for i, warning in enumerate(validation_warnings, 1):
                    st.markdown(f"**{i}. {warning['title']}**")
                    st.warning(warning['message'])
                    
                    # Show affected rows
                    with st.container():
                        st.markdown(f"**📋 Affected rows ({warning['count']} total):**")
                        for detail in warning['details']:
                            st.text(f"  • {detail}")
                        if warning['count'] > 10:
                            st.caption(f"  ... and {warning['count'] - 10} more rows")
                    st.markdown("")
    else:
        st.success("✅ **All data validation passed!** Ready to download or generate CVs.", icon="✅")

    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("#### 📄 Download Excel")
        
        # Disable download if there are critical errors
        can_download = len(validation_errors) == 0
        
        if not can_download:
            st.error("⚠️ Cannot download: Fix critical issues above")
            st.button(
                "⬇️ Download Excel",
                disabled=True,
                use_container_width=True,
                help="Fix all critical validation errors first"
            )
        else:
            # Prepare Excel download with all changes
            buf = io.BytesIO()
            df_download = st.session_state.df_personnel.copy()
            # Remove Assigned Role if exists
            if "Assigned Role" in df_download.columns:
                df_download = df_download.drop(columns=["Assigned Role"])
            # Ensure To is "Present" and YOE is integer
            df_download["To"] = "Present"
            if "Years of Experience" in df_download.columns:
                df_download["Years of Experience"] = df_download["Years of Experience"].apply(lambda x: int(float(x)) if pd.notna(x) else 0)
            df_download.to_excel(buf, index=False)
            buf.seek(0)
            
            st.download_button(
                "⬇️ Download Excel",
                buf,
                file_name="Personnel_download.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        
        st.caption("Downloads current personnel data with all your changes")

    with col2:
        st.markdown("#### 📝 Bulk CV Generation")
        
        # Disable CV generation if there are critical errors
        can_generate = len(validation_errors) == 0
        
        if not can_generate:
            st.error("⚠️ Cannot generate CVs: Fix critical issues above")
            st.button(
                "🚀 Generate CVs",
                disabled=True,
                type="primary",
                use_container_width=True,
                help="Fix all critical validation errors first"
            )
            st.caption("CVs cannot be generated with missing job titles")
        elif st.button("🚀 Generate CVs", type="primary", use_container_width=True):
            # Check if project info is available
            if st.session_state.df_project_info is None:
                st.error("❌ Project info not loaded. Please load project info in Step 1.")
            else:
                try:
                    # Show progress
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    status_text.text("⏳ Preparing data...")
                    progress_bar.progress(10)
                    
                    # Use current personnel data from session
                    df_personnel = st.session_state.df_personnel.copy()
                    df_projects = st.session_state.df_project_info.copy()
                    
                    status_text.text("⏳ Generating CVs...")
                    progress_bar.progress(30)
                    
                    # Run bulk generator
                    run_bulk_generator(
                        personnel_df=df_personnel,
                        project_info_df=df_projects,
                        out_docx=OUTPUT_DOCX
                    )
                    
                    progress_bar.progress(80)
                    status_text.text("⏳ Preparing download...")
                    
                    # Read the generated file
                    with open(OUTPUT_DOCX, "rb") as f:
                        docx_bytes = f.read()
                    
                    progress_bar.progress(100)
                    status_text.text("✅ Generation complete!")
                    
                    # Provide download button
                    st.download_button(
                        "⬇️ Download Employees_CV.docx",
                        docx_bytes,
                        file_name="Employees_CV.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
                    # Clean up - delete the file after download
                    if os.path.exists(OUTPUT_DOCX):
                        os.remove(OUTPUT_DOCX)
                    
                    st.success(f"✅ Generated CVs for {len(df_personnel)} personnel with random project assignments!")
                    
                except Exception as e:
                    st.error(f"❌ CV generation failed: {e}")
                    import traceback
                    st.code(traceback.format_exc())
        
        if can_generate:
            st.caption("Generates Word document with CVs for all personnel")
